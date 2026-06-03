#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
From online video fetch audio to produce script txt and answer docx

从在线视频链接提取音频，转写为文本，可选 LLM 纠错与问答，输出结构化 docx。
依赖联网依照用户选择自动下载构建，有 python 即可跑，无需手动为构建头疼。

⚠️ 仅下载本人拥有完整版权的原创内容，或取得授权的内容。
"""

import sys
import os
import re
import json
import time
import subprocess
import argparse
import textwrap
import shutil
import urllib.request
import urllib.error
from pathlib import Path
from datetime import timedelta

# ── 会话级"保持此操作"标志 ──────────────────────────
_keep_cache_action = None    # None/True(刷新)/False(用缓存)
_keep_skip_audio = None      # None/True(跳过)/False(重新下载)

# ─── 注册 NVIDIA CUDA DLL 路径（CTranslate2 需要 cublas64_12.dll 等）────
import site
for _sp in site.getsitepackages():
    for _pkg in ("cublas", "cudnn"):
        _lib = Path(_sp) / "nvidia" / _pkg / "lib"
        if _lib.exists():
            os.add_dll_directory(str(_lib))
            os.environ["PATH"] = str(_lib) + os.pathsep + os.environ.get("PATH", "")

# ─── 颜色输出（Windows 兼容）────────────────────────────
try:
    from colorama import init
    init(convert=True)
    HAS_COLORAMA = True
except ImportError:
    HAS_COLORAMA = False

def c(text, color):
    """带颜色的字符串"""
    if not HAS_COLORAMA:
        return text
    colors = {
        "RED":     "\033[91m",
        "GREEN":   "\033[92m",
        "YELLOW":  "\033[93m",
        "BLUE":    "\033[94m",
        "MAGENTA": "\033[95m",
        "CYAN":    "\033[96m",
        "WHITE":   "\033[97m",
        "RESET":   "\033[0m",
    }
    return f"{colors.get(color, '')}{text}{colors['RESET']}"

# ─── 路径常量 ───────────────────────────────────────────
SCRIPT_DIR    = Path(__file__).resolve().parent
AUDIO_DIR     = SCRIPT_DIR / "audio"
TRANSCRIPT_DIR = SCRIPT_DIR / "transcripts"
OUTPUT_DIR    = SCRIPT_DIR / "output"

def get_site_packages():
    """获取当前 Python 的 site-packages 路径（用于显示给用户）"""
    import sysconfig
    return sysconfig.get_paths()['purelib']

def get_free_disk_space_mb(path):
    """获取路径所在磁盘的剩余空间（单位：MB）"""
    try:
        total, used, free = shutil.disk_usage(path)
        return free // (1024 * 1024)
    except Exception:
        return None

TEMPLATE_HELP = """
====================================================================
                    模板格式说明
====================================================================

本工具支持两种模板类型。

【自然语言模板】
  直接用 标签名： 或 标签名 的格式列出想要提取的内容。
  工具会在首次使用时自动转换为标识符模板，供你检查和修改后再执行。

【标识符模板】
  使用标记语法精确控制输出格式：

  /.\保留字/.\：（【{字段名}】）
      → /.\保留字/.\： 原样保留在输出中（去除 /.\ 和 \./ 标记后）
      → （【{字段名}】） 替换为 AI 提取的内容
      → 最终输出示例：辩题：人工智能利大于弊

  特殊格式（无保留字）：
  （【{transcript}】）
      → 替换为完整转写文本（受 --segments 参数影响）

  预定义字段（从视频元数据获取）：
      index, title, url, duration, duration_fmt,
      upload_date, view_count, danmaku_count, bvid

  自定义字段（AI 语义提取，从转写文本中理解内容）：
      如 {辩题} 会理解讲话内容，提取辩题相关论述
      不要求转写文本中出现"辩题"这个字眼

====================================================================
"""

# ─── 必须依赖（不装就退出）──────────────────────────────
DEP_MAP_BASE = {
    "yt_dlp":   "yt-dlp",
    "docx":     "python-docx",
    "certifi":  "certifi",
    "httpx":    "httpx",
    "opencc":   "opencc-python-reimplemented",
    "colorama": "colorama",
}
# Whisper 相关（第2步按需安装）
DEP_MAP_WHISPER = {
    "faster_whisper": "faster-whisper",
}
# ─── LLM 模型选项（语义提取用）─────────────────────────
LLM_MODELS = {
    "qwen2.5-3b": {
        "repo": "qwen/Qwen2.5-3B-Instruct-GGUF",
        "file": "qwen2.5-3b-instruct-q4_k_m.gguf",
        "size_mb": 1800,
        "desc": "Qwen2.5-3B Q4_K_M  — 基准选择，CPU 也能跑，需 ~4GB 内存",
    },
    "qwen2.5-7b": {
        "repo": "qwen/Qwen2.5-7B-Instruct-GGUF",
        "file": "qwen2.5-7b-instruct-q4_k_m.gguf",
        "size_mb": 4700,
        "desc": "Qwen2.5-7B Q4_K_M  — 效果更好，推荐 CUDA 用户，需 ~8GB 显存",
    },
}
LLM_DEFAULT = "qwen2.5-3b"
# llama-cpp-python server 地址（脚本内启动，无需手动运行）
LLAMA_SERVER_URL = "http://127.0.0.1:8000/v1"
LLAMA_SERVER_PORT = 8000
LLAMA_SERVER_N_CTX = 8192  # server context，覆盖 chunk + prompt + 回复
LLAMA_MODEL_ID = None  # 启动后从 /v1/models 获取真实 model id


def _scan_local_gguf():
    """扫描 models/ 目录下的所有 .gguf 文件，返回可用模型字典。

    与 LLM_MODELS 不冲突：已有预定义的用预定义信息，新增的构造 [本地] 条目。
    返回: {key: {file, size_mb, desc, repo}}，repo 为 None 表示纯本地（不可下载）。
    """
    model_dir = SCRIPT_DIR / "models"
    if not model_dir.exists():
        return {}
    # 收集预定义中的文件名，避免重复显示
    known_files = {info["file"] for info in LLM_MODELS.values()}
    result = {}
    for gguf in model_dir.glob("*.gguf"):
        name = gguf.name
        if name in known_files:
            continue  # 已在 LLM_MODELS 中
        stem = gguf.stem
        size_mb = int(gguf.stat().st_size / (1024 * 1024))
        result[stem] = {
            "file": name,
            "size_mb": size_mb,
            "desc": f"[本地] {stem}  (~{size_mb} MB)",
            "repo": None,
        }
    return result


def _all_llm_models():
    """合并预定义 + 本地发现的所有 LLM 模型"""
    merged = dict(LLM_MODELS)
    merged.update(_scan_local_gguf())
    return merged


def _resolve_llm_key(user_input, default=LLM_DEFAULT):
    """大小写不敏感的模糊匹配：用户输入 → 精确 model key。

    - 空输入 → 返回 default
    - 唯一子串命中 → 自动选
    - 多个命中 → 列出候选项，提示再输入
    - 零命中 → 返回 default
    """
    all_models = _all_llm_models()
    if not user_input:
        return default
    user_input = user_input.strip().lower()
    lowered = {k.lower(): k for k in all_models}
    matches = [orig for low, orig in lowered.items() if user_input in low]
    if len(matches) == 1:
        return matches[0]
    elif len(matches) > 1:
        print(c(f"  匹配到多个模型，请指定：", "YELLOW"))
        for m in matches:
            print(c(f"    {m}", "WHITE"))
        retry = input(c("  重新输入: ", "CYAN")).strip().lower()
        return _resolve_llm_key(retry, default)
    # 零命中
    print(c(f"  [!] '{user_input}' 不是有效选项", "YELLOW"))
    return None


def _get_llm_info(model_key):
    """获取指定 LLM 模型的信息。先查预定义，再查本地发现，都找不到返回默认。"""
    all_models = _all_llm_models()
    return all_models.get(model_key, LLM_MODELS[LLM_DEFAULT])

WHISPER_DEFAULT = "tiny"
WHISPER_SIZE_MB = {
    "tiny":      75,
    "tiny.en":   75,
    "base":      145,
    "base.en":   145,
    "small":     488,
    "small.en":  488,
    "medium":    1536,
    "medium.en": 1536,
    "large-v3":  3072,
}

# ─── llama.cpp 预编译二进制（替代 llama-cpp-python，无需编译）────
LLAMA_CPP_RELEASE = "b9374"  # llama.cpp 发布版本 tag


# ═══════════════════════════════════════════════════
#  Section 1：依赖管理（自动检测硬件，按需安装）
# ═══════════════════════════════════════════════════

def detect_gpu():
    """
    统一 GPU 检测。先 nvidia-smi，再测 torch CUDA 状态，最后测 ctranslate2。
    ctranslate2 检测仅当 cublas/cudnn DLL 已安装时才有效。
    返回：(has_hw: bool, gpu_name: str, torch_cuda_ok: bool, ctranslate2_cuda_ok: bool)
    """
    has_hw, gpu_name = False, ""
    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=name", "--format=csv,noheader"],
            capture_output=True, text=True, timeout=8,
            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0,
        )
        if result.returncode == 0 and result.stdout.strip():
            has_hw = True
            gpu_name = result.stdout.strip().split("\n")[0].strip()
    except Exception:
        pass

    torch_cuda_ok = False
    try:
        import torch
        torch_cuda_ok = torch.cuda.is_available()
    except ImportError:
        pass

    ctranslate2_cuda_ok = False
    try:
        import ctranslate2
        if ctranslate2.get_cuda_device_count() > 0:
            ctranslate2_cuda_ok = True
    except Exception:
        pass

    return has_hw, gpu_name, torch_cuda_ok, ctranslate2_cuda_ok


def estimate_size(pkg_name, has_cuda=False):
    """估算 pip 包安装后占用空间（MB），仅供参考"""
    sizes = {
        "yt-dlp":            15,
        "faster-whisper":   200,
        "python-docx":       10,
        "certifi":            5,
        "httpx":             10,
        "opencc-python-reimplemented": 5,
        "colorama":           2,
    }
    if has_cuda:
        sizes["torch"]       = 2500
        sizes["torchaudio"]  = 2000
        sizes["nvidia-cublas-cu12"] = 300
        sizes["nvidia-cudnn-cu12"]  = 500
    return sizes.get(pkg_name, 50)

def check_deps(dep_map):
    """检查依赖是否已安装，返回缺失列表 [(模块名, pip包名), ...]"""
    missing = []
    for mod, pkg in dep_map.items():
        try:
            __import__(mod)
        except ImportError:
            missing.append((mod, pkg))
    return missing

def _run_pip(packages, extra_flags=None, force_reinstall=False, index_url=None):
    """执行 pip install，失败返回错误信息（成功返回 None）"""
    if index_url:
        pip_cmd = [sys.executable, "-m", "pip", "install",
                   "--force-reinstall" if force_reinstall else "--upgrade",
                   "--no-cache-dir",
                   "-i", index_url]
    else:
        pip_cmd = [sys.executable, "-m", "pip", "install",
                   "--force-reinstall" if force_reinstall else "--upgrade",
                   "--no-cache-dir",
                   "-i", "https://pypi.tuna.tsinghua.edu.cn/simple",
                   "--trusted-host", "pypi.tuna.tsinghua.edu.cn"]
    if extra_flags:
        pip_cmd.extend(extra_flags)
    pip_cmd.extend(packages)

    display = " ".join(pip_cmd)
    if len(display) > 100:
        display = display[:100] + "..."
    print(c(f"    执行：{display}", "WHITE"))

    try:
        subprocess.run(pip_cmd, check=True)
        return None
    except subprocess.CalledProcessError as e:
        return str(e)



def _check_whisper_cache(model_size):
    """检查 Whisper 模型是否已在 HF 缓存中"""
    base = Path.home() / ".cache" / "huggingface" / "hub" / f"models--Systran--faster-whisper-{model_size}"
    snapshots = base / "snapshots"
    if not snapshots.exists():
        return False
    for snap in snapshots.iterdir():
        if snap.is_dir() and list(snap.glob("*")):
            return True
    return False


def ensure_base_deps():
    """第1步：必须依赖（不装就退出）。返回 True 表示已就绪。"""
    dep_map = dict(DEP_MAP_BASE)
    missing = check_deps(dep_map)
    if not missing:
        print(c("  [OK] 基础依赖已齐全", "GREEN"))
        return True

    total = sum(estimate_size(pkg, False) for _, pkg in missing)
    print()
    print(c("=" * 52, "YELLOW"))
    print(c("  【必须依赖】以下包缺一不可", "YELLOW"))
    print(c("=" * 52, "YELLOW"))
    print()
    for mod, pkg in missing:
        mb = estimate_size(pkg, False)
        print(c(f"    * {pkg:<25} ~ {mb} MB", "CYAN"))
    print(c(f"  预计 ~{total} MB", "YELLOW"))
    print()

    ans = input(c("  是否安装？(y/n): ", "YELLOW")).strip().lower()
    if ans not in ("y", "yes", "shi", "hao", "ok"):
        print(c("\n  [X] 基础依赖不可跳过，退出。", "RED"))
        return False

    err = _run_pip([pkg for _, pkg in missing])
    if err:
        print(c(f"  [X] 安装失败: {err}", "RED"))
        return False
    print(c("  [OK] 基础依赖安装完成", "GREEN"))
    return True


def ensure_whisper_setup(model_size, use_cuda, gpu_name=""):
    """第2步：Whisper 依赖 + 模型。返回 True 表示已就绪。"""
    missing = check_deps(DEP_MAP_WHISPER)
    whisper_cached = _check_whisper_cache(model_size)

    # 检测 CUDA DLL：先问 ctranslate2 能不能用 CUDA（系统 CUDA 或 pip 包都能通过）
    cublas_needed = False
    if use_cuda:
        ctrans_cuda_ok = False
        try:
            import ctranslate2
            if ctranslate2.get_cuda_device_count() > 0:
                ctrans_cuda_ok = True
        except Exception:
            pass

        if ctrans_cuda_ok:
            cublas_needed = False
        else:
            # ctranslate2 不能用 CUDA，检查 pip 包 DLL 是否已存在
            cublas_needed = True
            import site
            for _sp in site.getsitepackages():
                if (Path(_sp) / "nvidia" / "cublas" / "lib" / "cublas64_12.dll").exists():
                    cublas_needed = False
                    break

    all_ok = not missing and whisper_cached and not cublas_needed
    if all_ok:
        print(c("  [OK] Whisper 依赖已齐全", "GREEN"))
        return True

    # 展示缺失清单
    total = 0
    items = []
    if missing:
        for mod, pkg in missing:
            mb = estimate_size(pkg, use_cuda)
            total += mb
            items.append(f"{pkg}  ~{mb} MB")
    if not whisper_cached:
        mb = WHISPER_SIZE_MB.get(model_size, 75)
        total += mb
        items.append(f"Whisper {model_size} 模型  ~{mb} MB")
    if cublas_needed:
        total += 800
        items.append("nvidia-cublas-cu12 + cudnn  ~800 MB")

    print()
    print(c("=" * 52, "YELLOW"))
    print(c("  【Whisper 依赖】需安装/下载：", "YELLOW"))
    print(c("=" * 52, "YELLOW"))
    print()
    if use_cuda:
        print(c(f"  模式：GPU 加速（{gpu_name}）", "GREEN"))
    else:
        print(c("  模式：CPU", "YELLOW"))
    print()
    for item in items:
        print(c(f"    * {item}", "CYAN"))
    print(c(f"  预计 ~{total} MB", "YELLOW"))
    print()

    ans = input(c("  是否安装？(y/n): ", "YELLOW")).strip().lower()
    if ans not in ("y", "yes", "shi", "hao", "ok"):
        print(c("  -> 跳过安装，可重新选择", "YELLOW"))
        return False

    # 安装 Python 包
    if missing:
        err = _run_pip([pkg for _, pkg in missing])
        if err:
            print(c(f"  [X] 安装失败: {err}", "RED"))
            return False

    # 下载 Whisper 模型
    if not whisper_cached:
        _ensure_whisper_cache(model_size)

    # 安装 cublas/cudnn
    if cublas_needed:
        print(c("  [*] 安装 ctranslate2 CUDA 运行时（cublas + cudnn）...", "CYAN"))
        cublas_err = _run_pip(["nvidia-cublas-cu12"], index_url="https://pypi.org/simple/")
        cudnn_err = _run_pip(["nvidia-cudnn-cu12"], index_url="https://pypi.org/simple/")
        if cublas_err or cudnn_err:
            print(c("  [X] GPU 加速依赖安装失败", "RED"))
            return False
        import site
        for _sp in site.getsitepackages():
            for _pkg in ("cublas", "cudnn"):
                _lib = Path(_sp) / "nvidia" / _pkg / "lib"
                if _lib.exists():
                    os.add_dll_directory(str(_lib))
                    os.environ["PATH"] = str(_lib) + os.pathsep + os.environ.get("PATH", "")

    print(c("  [OK] Whisper 依赖安装完成", "GREEN"))
    return True


def ensure_llm_setup(model_key, use_cuda):
    """第3步：LLM 模型 + llama.cpp 二进制。返回 True 表示已就绪。"""
    llm_status, llm_path = check_llm_model(model_key)
    llm_exists = (llm_status == "ok")
    llm_info = _get_llm_info(model_key)
    binary_exists, binary_path = check_llama_binary(use_cuda)

    if llm_exists and binary_exists:
        print(c("  [OK] LLM 依赖已齐全", "GREEN"))
        return True

    total = 0
    items = []
    if not llm_exists:
        total += llm_info["size_mb"]
        items.append(f"{llm_info['file']}  ~{llm_info['size_mb']} MB")
    if not binary_exists:
        total += 50
        variant = "CUDA" if use_cuda else "CPU"
        items.append(f"llama.cpp 二进制（{variant}版）  ~50 MB")

    print()
    print(c("=" * 52, "YELLOW"))
    print(c("  【LLM 依赖】需下载：", "YELLOW"))
    print(c("=" * 52, "YELLOW"))
    print()
    print(c(f"  llama.cpp: {'CUDA 版' if use_cuda else 'CPU 版'}", "YELLOW"))
    print()
    for item in items:
        print(c(f"    * {item}", "CYAN"))
    print(c(f"  预计 ~{total} MB", "YELLOW"))
    print()

    ans = input(c("  是否下载？(y/n): ", "YELLOW")).strip().lower()
    if ans not in ("y", "yes", "shi", "hao", "ok"):
        print(c("  -> 跳过下载，可重新选择", "YELLOW"))
        return False

    # 下载 LLM 模型
    if not llm_exists:
        result = ensure_llm_model(model_key)
        if result is None:
            print(c("  [X] LLM 模型下载失败", "RED"))
            return False
    else:
        print(c("  [OK] LLM 模型已存在", "GREEN"))

    # 下载 llama.cpp 二进制
    if not binary_exists:
        result = ensure_llama_binary(use_cuda)
        if result is None:
            print(c("  [X] llama.cpp 二进制下载失败", "RED"))
            return False
    else:
        print(c("  [OK] llama.cpp 二进制已就绪", "GREEN"))

    print(c("  [OK] LLM 依赖安装完成", "GREEN"))
    return True


def fix_ssl():
    """修复 SSL 证书问题（Windows 常见）"""
    try:
        import certifi
        os.environ["SSL_CERT_FILE"] = certifi.where()
        os.environ["REQUESTS_CA_BUNDLE"] = certifi.where()
    except ImportError:
        pass
    # monkey-patch httpx 使用 certifi
    try:
        import httpx
        import certifi
        orig_init = httpx.Client.__init__
        def patched_init(self, *args, **kwargs):
            if "verify" not in kwargs:
                kwargs["verify"] = certifi.where()
            return orig_init(self, *args, **kwargs)
        httpx.Client.__init__ = patched_init
    except Exception:
        pass


# ═══════════════════════════════════════════════════
#  Section 2：视频合集获取
# ═══════════════════════════════════════════════════

def _video_list_cache_path(url):
    """根据 URL 生成缓存文件路径"""
    import hashlib
    key = hashlib.md5(url.encode()).hexdigest()[:12]
    return SCRIPT_DIR / "meta_cache" / f"{key}.json"


def _try_resolve_collection(info, url):
    """
    当 yt-dlp 返回单视频时，调用 Bilibili API 检测 ugc_season 合集，
    若有则用合集 URL 重新拉取全部视频。
    成功返回 entries 列表，失败返回 None（打印提示，不中断流程）。
    """
    bvid = info.get("id") or info.get("bvid", "")
    if not bvid.startswith("BV"):
        return None

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Referer": "https://www.bilibili.com/",
    }

    # 第1步：调用 B站 API 获取 ugc_season 信息
    try:
        api_url = f"https://api.bilibili.com/x/web-interface/view?bvid={bvid}"
        req = urllib.request.Request(api_url, headers=headers)
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read())
    except Exception as e:
        print(c(f"  [!] 调用 B站 API 失败: {e}", "YELLOW"))
        return None

    if not isinstance(data, dict) or data.get("code") != 0:
        print(c("  [!] B站 API 返回异常", "YELLOW"))
        return None

    ugc_season = data.get("data", {}).get("ugc_season")
    if not ugc_season:
        print(c("  [!] 未检测到合集（API 无 ugc_season 字段），按单视频处理", "YELLOW"))
        return None

    sid = ugc_season["id"]
    mid = ugc_season["mid"]
    title = ugc_season.get("title", "")
    print(c(f"  [~] 检测到合集「{title}」(sid={sid})，正在拉取全部视频...", "YELLOW"))

    # 第2步：用 yt-dlp 拉取合集所有视频
    collection_url = (
        f"https://space.bilibili.com/{mid}/channel/collectiondetail?sid={sid}"
    )

    fix_ssl()
    import yt_dlp

    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "forcejson": True,
        "http_headers": headers,
    }

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            col_info = ydl.extract_info(collection_url, download=False)
        entries = col_info.get("entries")
        if entries is not None:
            count = (
                len(entries)
                if hasattr(entries, "__len__")
                else sum(1 for _ in entries)
            )
            print(c(f"  [OK] 合集共 {count} 个视频", "GREEN"))
            return entries
        print(c("  [!] 合集 URL 返回异常（无 entries）", "YELLOW"))
        return None
    except Exception as e:
        print(c(f"  [!] 合集拉取失败: {e}", "YELLOW"))
        return None


def fetch_video_list(url, force_refresh=False):
    """
    获取合集/视频列表（带重试 + 本地缓存）。
    缓存命中则直接返回，避免重复获取。
    返回 [{"index": int, "bvid": str, "title": str, "duration": int}, ...]
    """
    cache_path = _video_list_cache_path(url)

    # 读缓存
    if not force_refresh and cache_path.exists():
        import json
        try:
            cached = json.loads(cache_path.read_text(encoding="utf-8"))
            if isinstance(cached, list) and len(cached) > 0:
                global _keep_cache_action
                skip_cache = False
                if _keep_cache_action is None:
                    print(c(f"  [缓存] 发现 {len(cached)} 个视频的元数据缓存", "YELLOW"))
                    ans = input(c("  是否重新获取？(y/n，默认 n): ", "CYAN")).strip().lower()
                    skip_cache = ans in ("y", "yes")
                    ans2 = input(c("  保持此选择？(y/n，默认 y): ", "CYAN")).strip().lower()
                    if ans2 in ("y", "yes", ""):
                        _keep_cache_action = skip_cache
                else:
                    skip_cache = _keep_cache_action

                if not skip_cache:
                    print(c(f"  [OK] 使用缓存（{len(cached)} 个视频）", "GREEN"))
                    return cached
                # 跳过缓存 → 继续往下拉取
                print(c("  -> 重新获取元数据", "YELLOW"))
        except Exception:
            pass  # 缓存损坏，重新获取

    fix_ssl()
    import yt_dlp

    print(c("  正在获取视频列表...", "CYAN"))

    # ── yt-dlp 配置 ────────────────────────────────────────
    # 不加 extract_flat，让 yt-dlp 完整拉取才能拿到 合集全部视频
    # 加常见请求头避免被拒绝
    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "forcejson": True,
        "http_headers": {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Referer": "https://www.bilibili.com/",
        },
    }

    videos = []
    max_retries = 3
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
        except Exception as e:
            last_err = e
            print(c(f"  [!] 第 {attempt} 次获取失败: {e}", "YELLOW"))
            if attempt < max_retries:
                print(c(f"      {3 * attempt} 秒后重试...", "YELLOW"))
                time.sleep(3 * attempt)
                continue
            print(c(f"  [X] {max_retries} 次重试均失败: {last_err}", "RED"))
            return []

        entries = info.get("entries")
        if entries is None:
            entries = _try_resolve_collection(info, url)
        if entries is None:
            entries = [info]
        for i, entry in enumerate(entries, 1):
            bvid = entry.get("bvid") or entry.get("id", "")
            # 清理：去 _pN 后缀（yt-dlp 内部序号），补齐 BV 前缀（id 字段不含）
            bvid = re.sub(r'_p\d+$', '', bvid)
            if not bvid.startswith("BV"):
                bvid = "BV" + bvid
            title = entry.get("title", f"视频{i}")
            duration = entry.get("duration") or 0
            duration_sec = int(duration)
            h, m = divmod(duration_sec, 3600)
            m, s = divmod(m, 60)
            if h > 0:
                duration_fmt = f"{h}:{m:02d}:{s:02d}"
            else:
                duration_fmt = f"{m}:{s:02d}"
            videos.append({
                "index":         i,
                "bvid":          bvid,
                "title":         title,
                "duration":      duration_sec,
                "duration_fmt":  duration_fmt,
                "upload_date":   entry.get("upload_date") or "",
                "view_count":    entry.get("view_count") or 0,
                "danmaku_count": entry.get("danmaku_count") or 0,
                "url":           f"https://www.bilibili.com/video/{bvid}",
            })
        break  # 成功，跳出重试循环

    print(c(f"  [OK] 共找到 {len(videos)} 个视频", "GREEN"))
    if len(videos) == 1:
        print(c("  [!] 仅返回 1 个视频 — 如果是合集链接，可能是API 波动", "YELLOW"))

    # 写入缓存
    if videos:
        import json
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        cache_path.write_text(json.dumps(videos, ensure_ascii=False, indent=2), encoding="utf-8")

    return videos


def download_audio(video, audio_dir):
    """用 yt-dlp 下载视频音频（m4a 格式，无需 ffmpeg）"""
    bvid = video["bvid"]
    out_path = audio_dir / f"{bvid}.m4a"

    if out_path.exists():
        global _keep_skip_audio
        skip_dl = True
        if _keep_skip_audio is None:
            ans = input(c(f"    {bvid}.m4a 已存在，是否跳过下载？(y/n，默认 y): ", "CYAN")).strip().lower()
            skip_dl = ans not in ("n", "no", "fou", "bu")
            ans2 = input(c("    保持此选择？(y/n，默认 y): ", "CYAN")).strip().lower()
            if ans2 in ("y", "yes", ""):
                _keep_skip_audio = skip_dl
        else:
            skip_dl = _keep_skip_audio

        if skip_dl:
            print(c(f"    [OK] 音频已存在，跳过下载", "GREEN"))
            return str(out_path)
        print(c(f"    -> 重新下载音频", "YELLOW"))

    print(c(f"    [↓] 下载音频...", "CYAN"))
    fix_ssl()

    import yt_dlp
    ydl_opts = {
        "format": "worstaudio/bestaudio",
        "outtmpl": str(audio_dir / "%(id)s.%(ext)s"),
        "quiet": True,
        "no_warnings": True,
        "http_headers": {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Referer": "https://www.bilibili.com/",
        },
    }

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([video["url"]])
        if out_path.exists():
            print(c(f"    [OK] 下载完成: {out_path.name}", "GREEN"))
            return str(out_path)
        else:
            # 尝试找其他格式
            for ext in ("webm", "mp4", "mp3"):
                p = audio_dir / f"{bvid}.{ext}"
                if p.exists():
                    print(c(f"    [OK] 下载完成: {p.name}", "GREEN"))
                    return str(p)
            print(c(f"    [X] 下载后未找到音频文件", "RED"))
            return None
    except Exception as e:
        print(c(f"    [X] 下载失败: {e}", "RED"))
        return None


def resolve_short_url(url, timeout=10):
    """
    用 yt-dlp Python API 解析短链重定向，返回最终 URL。
    """
    try:
        from yt_dlp import YoutubeDL
        ydl_opts = {
            "quiet": True,
            "no_warnings": True,
            "extract_flat": True,
            "http_headers": {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0.0.0 Safari/537.36"
                ),
                "Referer": "https://www.bilibili.com/",
            },
        }
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            final = info.get("webpage_url") or url
            if final and "bilibili.com" in final and final != url:
                return final
    except Exception as e:
        print(c(f"  [!] 短链解析失败: {e}", "YELLOW"))
    return url

def extract_url(text):
    """从含标题的粘贴文本中提取网址，不限定平台"""
    idx = text.find("http")
    if idx == -1:
        return ""
    url = text[idx:].split()[0]
    url = url.rstrip("）】」'),>\"")
    # 如果是短链，尝试解析重定向
    if "b23.tv" in url or "bili.im" in url or "b23.com" in url:
        resolved = resolve_short_url(url)
        if resolved != url:
            print(f"  短链解析：{url}  →  {resolved}")
            url = resolved
    return url


def parse_segments(seg_str, duration_sec):
    """
    解析时间段参数，返回 [(start, end), ...] 列表（秒）。
    支持格式：
        首3分          -> [(0, 180)]
        尾2分          -> [(duration-120, duration)]
        1:30-5:45     -> [(90, 345)]
        1:30:00-2:15:00 -> [(5400, 8100)]
        0-3,75-80     -> [(0, 180), (4500, 4800)]
        空             -> [(0, duration_sec)]（全部）
    边界处理（统一后处理）：
        - start < 0       -> 截断到 0
        - end > duration   -> 截断到 duration
        - start > end      -> 交换（理解为 n,m 顺序无关）
        - start >= end     -> 跳过（空段）
    """
    if not seg_str or not seg_str.strip():
        return [(0, duration_sec)]

    ranges = []
    parts = re.split(r"[,\uff0c]", seg_str)  # 支持中英文逗号

    for part in parts:
        part = part.strip()
        if not part:
            continue

        s = e = None

        # "首X分"（首/头/初/始）
        m = re.match(r"[首头初始]\s*(\d+(?:\.\d+)?)\s*分?", part)
        if m:
            mins = float(m.group(1))
            s, e = 0, int(mins * 60)

        # "尾X分"（尾/末）
        if s is None:
            m = re.match(r"[尾末]\s*(\d+(?:\.\d+)?)\s*分?", part)
            if m:
                mins = float(m.group(1))
                s = max(0, duration_sec - int(mins * 60))
                e = duration_sec

        # "1:30-5:45" 或 "1:30:00-5:45:00"
        if s is None:
            m = re.match(
                r"(\d+):(\d+)(?::(\d+))?\s*[-~\u2014到]\s*(\d+):(\d+)(?::(\d+))?",
                part
            )
            if m:
                if m.group(3) is not None:
                    s = int(m.group(1))*3600 + int(m.group(2))*60 + int(m.group(3))
                    e = int(m.group(4))*3600 + int(m.group(5))*60 + int(m.group(6))
                else:
                    s = int(m.group(1))*60 + int(m.group(2))
                    e = int(m.group(4))*60 + int(m.group(5))

        # "0-3" （分钟）
        if s is None:
            m = re.match(r"(\d+(?:\.\d+)?)\s*[-~\u2014到]\s*(\d+(?:\.\d+)?)", part)
            if m:
                s = int(float(m.group(1)) * 60)
                e = int(float(m.group(2)) * 60)

        if s is None:
            print(c(f"    [!] 无法解析时间段: {part}", "YELLOW"))
            continue

        # ── 统一后处理 ──────────────────────────────
        # 1. 截断到合法边界
        s = max(0, s)
        e = min(duration_sec, e)

        # 2. start > end -> 交换
        if s > e:
            s, e = e, s
            print(c(f"    [~] 时间段 {part} 起止顺序已自动调整", "YELLOW"))

        # 3. 空段 -> 跳过
        if s >= e:
            print(c(f"    [!] 时间段 {part} 为空，已跳过", "YELLOW"))
            continue

        ranges.append((s, e))

    return ranges if ranges else [(0, duration_sec)]

def build_initial_prompt(labels, video_meta=None):
    """从模板标签 + 视频元数据构建 Whisper initial_prompt，引导偏向正确词汇"""
    parts = []
    # 来源 B：模板字段标签（准确术语）
    if labels:
        kw = " ".join(dict.fromkeys(labels.values()))
        if kw:
            parts.append(kw)
    # 来源 A：视频标题（已知正确的背景信息）
    if video_meta:
        title = video_meta.get("title", "")
        if title:
            parts.append(title)
    return " ".join(parts) if parts else None


def transcribe_audio(audio_path, model, seg_ranges, language, initial_prompt=None):
    """转写音频文件，一次解码，按 seg_ranges 过滤"""
    print(c(f"    [~] 转写中...", "CYAN"))
    if initial_prompt:
        print(c(f"    引导词: {initial_prompt}", "YELLOW"))

    # 一次解码整个音频，tqdm 自动显示进度条
    segments, _ = model.transcribe(
        audio_path,
        language=language,
        initial_prompt=initial_prompt,
        vad_filter=True,
        word_timestamps=False,
        log_progress=True,
    )

    # 按时间段过滤
    all_segments = []
    for seg in segments:
        seg_start = seg.start
        seg_end = seg.end
        for start_sec, end_sec in seg_ranges:
            if seg_end >= start_sec and seg_start <= end_sec:
                all_segments.append(seg.text.strip())
                break

    transcript = " ".join(all_segments)
    print(c(f"    [OK] 转写完成 ({len(transcript)} 字)", "GREEN"))
    return transcript


def fix_transcript_chinese(text):
    """繁体→简体 + VTT 标记清理"""
    try:
        from opencc import OpenCC
        cc = OpenCC("t2s")
        text = cc.convert(text)
    except Exception:
        pass  # opencc 不可用时跳过
    # 清理 VTT 标签残留
    text = re.sub(r"<[^>]*>", "", text)
    text = re.sub(r"\\[a-z]+\s*\{[^}]*\}", "", text)
    return text


# ═══════════════════════════════════════════════════
#  llama.cpp server（HTTP API，替代 subprocess）
# ═══════════════════════════════════════════════════

def _start_llama_server(model_path, use_cuda):
    """启动 llama.cpp server (llama-server.exe)，等待就绪后返回 Popen 对象。同时设置 LLAMA_MODEL_ID。"""
    global LLAMA_MODEL_ID
    subdir = _llama_subdir(use_cuda)
    server_exe = str(subdir / "llama-server.exe")
    if not Path(server_exe).exists():
        print(c(f"  [X] 未找到 llama-server.exe ({subdir})", "RED"))
        return None

    cmd = [server_exe,
           "--model", model_path,
           "--host", "127.0.0.1",
           "--port", str(LLAMA_SERVER_PORT),
           "--ctx-size", str(LLAMA_SERVER_N_CTX),
           "--parallel", "1",
           ]
    if use_cuda:
        cmd += ["--n-gpu-layers", "99"]
        gpu_note = " +GPU"
    else:
        gpu_note = ""
    proc = subprocess.Popen(
        cmd,
        cwd=str(subdir),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    print(c(f"  [~] 启动 LLM server (ctx={LLAMA_SERVER_N_CTX}{gpu_note})...", "CYAN"))
    for _ in range(120):
        if proc.poll() is not None:
            print(c("  [X] LLM server 启动失败（进程已退出）", "RED"))
            return None
        try:
            resp = urllib.request.urlopen(f"{LLAMA_SERVER_URL}/models", timeout=2)
            data = json.loads(resp.read())
            models = data.get("data", [])
            if models:
                LLAMA_MODEL_ID = models[0].get("id", "unknown")
                print(c(f"  [OK] LLM server 就绪 (model: {LLAMA_MODEL_ID})", "GREEN"))
            else:
                print(c("  [OK] LLM server 就绪", "GREEN"))
            return proc
        except Exception:
            time.sleep(1)
    print(c("  [X] LLM server 启动超时", "RED"))
    proc.kill()
    proc.wait()
    return None


def _stop_llama_server(proc):
    """关闭 LLM server"""
    if proc and proc.poll() is None:
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()
        print(c("  [OK] LLM server 已关闭", "GREEN"))


def _llm_chat(messages, max_tokens=None, temperature=0.1):
    """调用本地 LLM server，一次 HTTP POST，不设客户端超时。max_tokens 不传则由 server 根据剩余 context 自动决定。"""
    global LLAMA_MODEL_ID
    body = {
        "model": LLAMA_MODEL_ID or "unknown",
        "messages": messages,
        "temperature": temperature,
    }
    if max_tokens is not None:
        body["max_tokens"] = max_tokens
    req = urllib.request.Request(
        f"{LLAMA_SERVER_URL}/chat/completions",
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req) as resp:
            data = json.loads(resp.read())
        return data["choices"][0]["message"]["content"]
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"HTTP {e.code}: {err_body}") from e


def correct_homophones_with_llm(transcript, video_meta, labels=None, bvid=None):
    """
    用 LLM 修正转写文本中的同音字错误（通过 HTTP API 调 llama server）。
    以视频元数据（标题等）+ 模板字段标签为准确背景信息。
    失败则返回原文。
    若提供 bvid，纠错结果同时写入 transcripts/{bvid}-纠错后.txt。
    """
    title = video_meta.get("title", "")

    # 构建背景上下文
    context_parts = []
    if title:
        context_parts.append(f"视频标题：{title}")
    description = video_meta.get("description", "")
    if description:
        context_parts.append(f"视频简介：{description[:200]}")
    if labels:
        label_terms = " ".join(dict.fromkeys(labels.values()))
        if label_terms:
            context_parts.append(f"关键术语（请保持这些词语的拼写正确）：{label_terms}")
    context = "\n".join(context_parts)

    # 等长硬切，每块 2000 字
    CHUNK_SIZE = 2000
    chunks = [transcript[i:i + CHUNK_SIZE] for i in range(0, len(transcript), CHUNK_SIZE)]

    print(c(f"    [~] LLM 同音字纠错中...({len(chunks)} 块)", "CYAN"))
    sys.stdout.flush()

    corrected_chunks = []
    for i, chunk in enumerate(chunks, 1):
        print(f"    [{i}/{len(chunks)}] LLM 同音字纠错中...", end="\r")
        sys.stdout.flush()
        try:
            raw = _llm_chat(
                messages=[
                    {"role": "system", "content": (
                        "你是语音转文字的校对助手。根据背景信息修正同音字错误"
                        "（如「变体」→「辩题」、「在见」→「再见」），"
                        "不要修改语义和句子结构，不要添加内容。仅输出修正后的本段文本。"
                    )},
                    {"role": "user", "content": (
                        f"背景信息：\n{context}\n\n"
                        f"以下语音转文字可能有同音字错误，请修正：\n\n{chunk}"
                    )},
                ],
            )
            raw = raw.strip()
            if len(raw) < len(chunk) * 0.3:
                print(c(f"    [!] 块 {i}/{len(chunks)} 输出异常，使用原文", "YELLOW"))
                corrected_chunks.append(chunk)
            else:
                corrected_chunks.append(raw)
        except Exception as e:
            print(c(f"    [!] 块 {i}/{len(chunks)} 失败: {e}，使用原文", "YELLOW"))
            corrected_chunks.append(chunk)
    print()

    result_text = "".join(corrected_chunks)
    print(c(f"    [OK] 纠错完成 ({len(transcript)}→{len(result_text)} 字)", "GREEN"))

    # 输出纠错后文件
    if bvid:
        corrected_path = TRANSCRIPT_DIR / f"{bvid}-纠错后.txt"
        corrected_path.write_text(result_text, encoding="utf-8")
        print(c(f"    [OK] 纠错后文本已保存: {corrected_path.name}", "GREEN"))

    return result_text




def _ensure_whisper_cache(model_size):
    """主动预下载 Whisper 模型到 HF 缓存（供 ensure_deps 调用）"""
    model_map = {
        "tiny":      "Systran/faster-whisper-tiny",
        "tiny.en":   "Systran/faster-whisper-tiny.en",
        "small":     "Systran/faster-whisper-small",
        "small.en":  "Systran/faster-whisper-small.en",
        "base":      "Systran/faster-whisper-base",
        "base.en":   "Systran/faster-whisper-base.en",
        "medium":    "Systran/faster-whisper-medium",
        "medium.en": "Systran/faster-whisper-medium.en",
        "large-v3":  "Systran/faster-whisper-large-v3",
    }
    repo_id = model_map.get(model_size, f"Systran/faster-whisper-{model_size}")
    try:
        from huggingface_hub import snapshot_download
        print(c(f"  [↓] 下载 Whisper {model_size}（约 {WHISPER_SIZE_MB.get(model_size, 75)} MB）...", "CYAN"))
        snapshot_download(
            repo_id,
            endpoint="https://hf-mirror.com",
            local_files_only=False,
        )
        print(c(f"  [OK] Whisper {model_size} 下载完成", "GREEN"))
        return True
    except Exception as e:
        print(c(f"  [!] Whisper {model_size} 预下载失败: {e}", "YELLOW"))
        print(c("    将在加载模型时自动重试。", "YELLOW"))
        return False


def _fix_whisper_snapshot(model_size):
    """修复 HuggingFace 缓存：Windows 下 blob 文件不会自动链接到 snapshots 目录"""
    model_map = {
        "tiny":      "Systran--faster-whisper-tiny",
        "tiny.en":   "Systran--faster-whisper-tiny.en",
        "small":     "Systran--faster-whisper-small",
        "small.en":  "Systran--faster-whisper-small.en",
        "base":      "Systran--faster-whisper-base",
        "base.en":   "Systran--faster-whisper-base.en",
        "medium":    "Systran--faster-whisper-medium",
        "medium.en": "Systran--faster-whisper-medium.en",
        "large-v3":  "Systran--faster-whisper-large-v3",
    }
    model_dir = model_map.get(model_size, f"Systran--faster-whisper-{model_size}")
    base = Path.home() / ".cache" / "huggingface" / "hub" / f"models--{model_dir}"
    blobs = base / "blobs"
    if not blobs.exists():
        return

    snap_dir = base / "snapshots"
    if not snap_dir.exists():
        return
    snapshots = sorted(snap_dir.iterdir())
    if not snapshots:
        return
    snap = snapshots[0]

    # 已有文件就不重复复制
    existing = list(snap.glob("*"))
    if len(existing) >= 5:
        return

    name_map = {}
    try:
        blob_hashes = os.listdir(str(blobs))
    except Exception:
        return

    for h in blob_hashes:
        hpath = blobs / h
        try:
            sz = hpath.stat().st_size
        except Exception:
            continue
        if sz > 70_000_000:
            name_map[h] = "model.bin"
        elif sz > 2_000_000:
            name_map[h] = "tokenizer.json"
        elif sz > 400_000:
            name_map[h] = "vocabulary.txt"
        elif sz > 1000:
            name_map[h] = "config.json"
        else:
            name_map[h] = "preprocessor_config.json"

    for h, fname in name_map.items():
        src = blobs / h
        dst = snap / fname
        if not dst.exists():
            try:
                shutil.copy2(str(src), str(dst))
            except Exception:
                pass


def load_whisper_model(model_size, device):
    """加载 faster-whisper 模型"""
    fix_ssl()
    # 设置 HuggingFace 国内镜像（必须在 import 之前，否则直连 hf.co 超时/被墙）
    os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
    from faster_whisper import WhisperModel

    compute_type = "float16" if device == "cuda" else "int8"

    # 修复 Windows 下 HuggingFace blob→snapshot 符号链接失败（必须在加载前）
    _fix_whisper_snapshot(model_size)

    cpu_threads = os.cpu_count() if device == "cpu" else 0
    thread_info = f", {cpu_threads} threads" if cpu_threads else ""
    print(c(f"  [模型] 加载 faster-whisper ({model_size}, {device}, {compute_type}{thread_info})...", "CYAN"))

    model = WhisperModel(
        model_size,
        device=device,
        compute_type=compute_type,
        cpu_threads=cpu_threads,
        num_workers=1,
        local_files_only=False,
    )

    print(c(f"  [模型] 加载完成", "GREEN"))
    return model


def _verify_gguf(path, expected_size_mb=None):
    """验证 GGUF 文件完整性：魔数 + 实际大小 ≥ 预期的 85%。"""
    try:
        with open(path, "rb") as f:
            if f.read(4) != b"GGUF":
                return False
        if expected_size_mb:
            actual_mb = os.path.getsize(path) / (1024 * 1024)
            if actual_mb < expected_size_mb * 0.85:
                return False
        return True
    except Exception:
        return False


def check_llm_model(model_key=LLM_DEFAULT):
    """检查 LLM GGUF 模型状态。返回 (status, path)：
    status: "ok" 完好 / "corrupt" 损坏 / "missing" 不存在
    """
    info = _get_llm_info(model_key)
    model_path = SCRIPT_DIR / "models" / info["file"]
    if not model_path.exists():
        return "missing", str(model_path)
    if _verify_gguf(str(model_path), info.get("size_mb")):
        return "ok", str(model_path)
    return "corrupt", str(model_path)


def ensure_llm_model(model_key=LLM_DEFAULT):
    """下载指定 LLM GGUF 模型（从 ModelScope，国内可用）。纯本地模型跳过下载。损坏文件自动删除重下。"""
    info = _get_llm_info(model_key)
    model_dir = SCRIPT_DIR / "models"
    model_path = model_dir / info["file"]

    if model_path.exists():
        if _verify_gguf(str(model_path), info.get("size_mb")):
            print(c(f"  [OK] LLM 模型已存在: {model_path.name}", "GREEN"))
            return str(model_path)
        else:
            print(c(f"  [!] 模型文件损坏，删除后重试下载...", "YELLOW"))
            model_path.unlink()

    # 纯本地模型（无下载源），且文件不存在
    if info.get("repo") is None:
        print(c(f"  [X] 模型 {info['file']} 不存在且无下载源，请手动放入 models/ 目录", "RED"))
        return None

    print(c(f"  [↓] 下载 LLM 模型 {info['file']} (~{info['size_mb']} MB)...", "CYAN"))
    model_dir.mkdir(parents=True, exist_ok=True)

    url = f"https://www.modelscope.cn/models/{info['repo']}/resolve/master/{info['file']}"

    try:
        from urllib.request import urlretrieve

        def _progress(block_num, block_size, total_size):
            if total_size <= 0:
                return
            done = block_num * block_size
            pct = min(100, done * 100 // total_size)
            print(f"\r    下载进度: {pct}%", end="", flush=True)

        urlretrieve(url, str(model_path), _progress)
        print()
        print(c(f"  [OK] LLM 模型下载完成: {model_path.name}", "GREEN"))
        return str(model_path)
    except Exception as e:
        print()
        print(c(f"  [X] 下载失败: {e}", "RED"))
        print(c(f"    请手动下载 {url} 到 {model_path}", "YELLOW"))
        return None


def _llama_subdir(use_cuda):
    """返回指定 GPU 状态对应的 llama.cpp 二进制子目录"""
    variant = "cuda" if use_cuda else "cpu"
    return SCRIPT_DIR / "tools" / "llama.cpp" / variant


def check_llama_binary(use_cuda):
    """检查指定 GPU 状态下的 llama.cpp 二进制是否已就绪。

    CPU 版：有 llama-server.exe 即就绪。
    CUDA 版：还需三个 CUDA 运行时 DLL（cudart64_12.dll / cublas64_12.dll / cublasLt64_12.dll）。
    """
    subdir = _llama_subdir(use_cuda)
    exe_path = subdir / "llama-server.exe"
    if not exe_path.exists():
        return False, None
    # CUDA 版本额外检查运行时 DLL
    if use_cuda:
        for dll in ["cudart64_12.dll", "cublas64_12.dll", "cublasLt64_12.dll"]:
            if not (subdir / dll).exists():
                return False, None
    return True, str(exe_path)


def ensure_llama_binary(use_cuda):
    """
    下载 llama.cpp 预编译二进制（从 GitHub Releases，国内需代理）。
    use_cuda=True → CUDA 版，False → 纯 CPU 版。
    两个版本分目录存放：tools/llama.cpp/cpu/ 和 tools/llama.cpp/cuda/。
    下载 zip → 解压到对应子目录 → 返回 llama-server.exe 路径，失败返回 None。
    """
    suffix = "cuda-12.4-x64" if use_cuda else "x64"
    variant = "cuda" if use_cuda else "cpu"
    zip_name = f"llama-{LLAMA_CPP_RELEASE}-bin-win-{suffix}.zip"
    download_url = (
        f"https://github.com/ggml-org/llama.cpp/releases/download/"
        f"{LLAMA_CPP_RELEASE}/{zip_name}"
    )
    tools_dir = _llama_subdir(use_cuda)
    zip_path = SCRIPT_DIR / "tools" / zip_name

    # 优先本地已有 zip（上次下载残留）
    if zip_path.exists():
        print(c(f"  [OK] 发现本地离线包: {zip_name}", "GREEN"))
    else:
        print(c(f"  [↓] 下载 llama.cpp 预编译二进制 (~50 MB)...", "CYAN"))
        print(c(f"     GitHub Releases 下载可能需要代理/VPN", "YELLOW"))
        print(c(f"     URL: {download_url}", "YELLOW"))
        tools_dir.parent.mkdir(parents=True, exist_ok=True)

        # 用 curl 下载（支持代理，有进度条）
        try:
            subprocess.run([
                "curl", "-L", "-#", "--ssl-no-revoke", "-o", str(zip_path),
                download_url
            ], check=True)
            print()
        except (subprocess.CalledProcessError, FileNotFoundError):
            # curl 不可用，回退到 urllib
            print(c("     curl 不可用，尝试 urllib（可能较慢）...", "YELLOW"))
            try:
                from urllib.request import urlretrieve
                urlretrieve(download_url, str(zip_path))
            except Exception as e:
                print(c(f"  [X] 下载失败: {e}", "RED"))
                print(c(f"    请手动下载 {download_url}", "YELLOW"))
                print(c(f"    解压到 {tools_dir}", "YELLOW"))
                return None

    # 解压到对应子目录
    print(c(f"  [↓] 解压到 {variant}/ ...", "CYAN"))
    import zipfile
    tools_dir.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(str(zip_path), "r") as zf:
            zf.extractall(str(tools_dir))
    except zipfile.BadZipFile:
        print(c("  [X] ZIP 文件损坏，删除后重试", "RED"))
        zip_path.unlink(missing_ok=True)
        return None

    # 删除主包 zip（节省空间）
    zip_path.unlink(missing_ok=True)

    # CUDA 补充包：cudart64_12.dll / cublas64_12.dll / cublasLt64_12.dll
    if use_cuda:
        cudart_name = "cudart-llama-bin-win-cuda-12.4-x64.zip"
        cudart_url = (
            f"https://github.com/ggml-org/llama.cpp/releases/download/"
            f"{LLAMA_CPP_RELEASE}/{cudart_name}"
        )
        cudart_path = SCRIPT_DIR / "tools" / cudart_name
        while True:
            if cudart_path.exists():
                print(c(f"  [OK] 发现本地 CUDA 补充包: {cudart_name}", "GREEN"))
            else:
                print(c(f"  [↓] 下载 CUDA 运行时补充包 (~10 MB)...", "CYAN"))
                try:
                    subprocess.run([
                        "curl", "-L", "-#", "--ssl-no-revoke", "-o", str(cudart_path),
                        cudart_url
                    ], check=True)
                    print()
                except (subprocess.CalledProcessError, FileNotFoundError):
                    print(c("     curl 不可用，尝试 urllib...", "YELLOW"))
                    try:
                        from urllib.request import urlretrieve
                        urlretrieve(cudart_url, str(cudart_path))
                    except Exception as e:
                        print(c(f"  [X] 补充包下载失败: {e}", "RED"))
                        cudart_path.unlink(missing_ok=True)
                        continue
            if not cudart_path.exists():
                continue
            print(c(f"  [↓] 解压补充包到 {variant}/ ...", "CYAN"))
            try:
                with zipfile.ZipFile(str(cudart_path), "r") as zf:
                    zf.extractall(str(tools_dir))
                cudart_path.unlink(missing_ok=True)
                print(c(f"  [OK] CUDA 运行时 DLL 就绪", "GREEN"))
                break
            except zipfile.BadZipFile:
                print(c("  [X] 补充包 ZIP 损坏，删除后重试", "RED"))
                cudart_path.unlink(missing_ok=True)
                continue

    # 直取子目录下的 llama-server.exe
    exe_path = tools_dir / "llama-server.exe"
    if exe_path.exists():
        print(c(f"  [OK] llama.cpp 二进制就绪: {exe_path.name}", "GREEN"))
        return str(exe_path)

    print(c("  [X] 解压后未找到 llama-server.exe，请检查", "RED"))
    return None


def _find_keyword_context(keyword, text, window=2, max_chars=3000):
    """
    在文本中搜索关键词，返回所有出现位置 ±window 句的上下文拼接。
    多次命中的区间如重叠则合并。总字数超过 max_chars 后截断。
    找不到返回 None。
    """
    sentences = re.split(r"(?<=[。！？\n])", text)
    sentences = [s for s in sentences if s.strip()]
    if not sentences:
        return None

    # 找出所有含关键词的句子索引
    hit_indices = [i for i, s in enumerate(sentences) if keyword in s]
    if not hit_indices:
        return None

    # 每个命中点取 ±window，合并重叠区间
    ranges = []
    for idx in hit_indices:
        lo = max(0, idx - window)
        hi = min(len(sentences) - 1, idx + window)
        ranges.append((lo, hi))

    # 合并重叠/相邻区间
    ranges.sort()
    merged = [ranges[0]]
    for lo, hi in ranges[1:]:
        prev_lo, prev_hi = merged[-1]
        if lo <= prev_hi + 1:  # 相邻也合并
            merged[-1] = (prev_lo, max(prev_hi, hi))
        else:
            merged.append((lo, hi))

    # 拼接，硬上限 max_chars（防关键词密集分布爆 context）
    parts = []
    total = 0
    for lo, hi in merged:
        chunk = "".join(sentences[lo:hi + 1])
        if total + len(chunk) > max_chars:
            break
        parts.append(chunk)
        total += len(chunk)
    return "\n---\n".join(parts) if parts else None


def extract_field_with_llm(field_name, transcript, label=None, video_meta=None):
    """
    用 LLM 从转写文本中语义提取字段值（通过 HTTP API 调 llama-cpp-python server）。
    优先关键词定位裁剪上下文，找不到再全量分段处理。
    field_name: 字段名
    label:      /.\标签/.\ 中的搜索描述文本
    video_meta: 视频元数据（作为领域背景帮助 LLM 判断）
    """
    target = label if label else field_name
    # 尝试关键词定位，裁剪上下文
    context_texts = None
    context = _find_keyword_context(target, transcript, window=2, max_chars=3000)
    if context:
        context_texts = [context]
    else:
        max_chunk = 2000
        sentences = re.split(r"(?<=[。！？\n])", transcript)
        chunks = []
        buf = ""
        for s in sentences:
            if len(buf) + len(s) > max_chunk and buf:
                chunks.append(buf)
                buf = s
            else:
                buf += s
        if buf:
            chunks.append(buf)
        context_texts = chunks if chunks else [transcript]

    for i, text in enumerate(context_texts):
        # ── 领域背景（system） ──
        sys_parts = ["你是一个文本提取助手，只输出提取结果，不要说废话。"]
        if video_meta:
            if video_meta.get("title"):
                sys_parts.append(f"领域背景：这个视频的标题是「{video_meta['title']}」。你对这个领域已有了解。")
            desc = video_meta.get("description", "")
            if desc and desc.strip():
                sys_parts.append(f"补充信息：{desc[:150]}")

        # ── 提取任务（user） ──
        user_text = (
            f"【以下是要检索的转写文本】\n{text}\n\n"
            f"【提取任务】\n"
            f"「{target}」是分类标签，你要提取的是 {field_name} 对应的具体内容。\n"
            f"示例：如问「辩题」，输出'人工智能利大于弊'，而不是找'辩题'这个词本身。\n"
            f"找不到则输出「（未找到）」。"
        )

        try:
            raw = _llm_chat(
                messages=[
                    {"role": "system", "content": "\n".join(sys_parts)},
                    {"role": "user", "content": user_text},
                ],
            )
            raw = raw.strip("。，：:  \n\r")
            if not raw or "(未找到)" in raw or "找不到" in raw or "未找到" in raw:
                continue
            return raw
        except Exception as e:
            print(c(f"    [!] LLM 提取 {field_name} 失败: {e}", "YELLOW"))
            continue

    return ""


# ═══════════════════════════════════════════════════
#  Section 4：模板解析 + 字段提取
# ═══════════════════════════════════════════════════

def _extract_labels_and_fields(raw):
    """
    同时提取 /.\\标签/.\ 和 （【{字段}】），严格成对匹配。
    返回：(fields: list, labels: dict{field_name: label_text})
    - /.\\标签/.\  → LLM 搜索目标
    - （【{字段}】）→ 填充位置 key
    同一行内两者同时存在 → 建立映射；填充位置依附搜索目标存在。
    """
    placeholders = re.findall(r"（【\{([^}]+)\}】）", raw)
    fields = [p.strip() for p in placeholders]

    labels = {}
    for line in raw.split("\n"):
        lbl_m = re.search(r"/\.\\(.+?)/\.\\", line)
        fld_m = re.search(r"（【\{([^}]+)\}】）", line)
        if lbl_m and fld_m:
            labels[fld_m.group(1).strip()] = lbl_m.group(1).strip()

    return fields, labels


def natural_to_identifier(text):
    """
    将自然语言模板转换为标识符模板。
    支持两种格式：
      标签名： 或 标签名:  → 转为 /.\标签名/\.：（【{标签名}】）
      整行就是一个标签名（无冒号）→ 转为 /.\标签名/\.：（【{标签名}】）
    """
    lines = text.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append(line)
            continue
        # 分隔线（纯标点重复行）原样保留，不转标识符
        if re.match(r'^[=\-\.\*#~\s]+$', stripped):
            result.append(line)
            continue
        # 已是标识符格式（有 /.\ 标记），原样保留
        if "/.\\" in stripped:
            result.append(line)
            continue

        # 匹配 "标签名：" 或 "标签名:" 格式（仅行首，跳过缩进）
        m = re.match(r'(\s*)([^\s:：]{1,30}?)\s*([：:])\s*(.*?)\s*$', line)
        if m:
            indent = m.group(1)
            label = m.group(2).strip()
            colon = m.group(3)
            rest = m.group(4)
            # 转换为新标识符格式：/.\label/.\： + （【{label}】）
            result.append(f"{indent}/.\\{label}/.\\{colon}（【{{{label}}}】）")
            continue

        # 整行就是一个标签名（无冒号，非空，非标记行）
        if re.match(r'^\s*.{1,30}\s*$', line) and not stripped.startswith(("（【", "}）")):
            indent = line[:len(line) - len(line.lstrip())]
            label = stripped
            result.append(f"{indent}/.\\{label}/.\\：（【{{{label}}}】）")
            continue

        result.append(line)
    return "\n".join(result)


def confirm_template(template_path):
    """
    模板确认阶段。
    - 自然语言模板（无 /.\ 标记）→ 转换 → 保存 → 展示 → 等用户修改确认
    - 标识符模板（有 /.\ 标记或 （【{…}】）标记）→ 展示字段 → 等用户确认
    返回 (fields, labels, template_raw) — 以确认后的标识符模板为准。
    """
    path = Path(template_path)

    # ── 文件不存在处理 ────────────────────────────────────
    if not path.exists():
        print(c(f"  [X] 模板文件不存在: {template_path}", "RED"))
        return [], {}, ""

    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        raw = "\n".join(p.text for p in doc.paragraphs)
    else:
        raw = path.read_text(encoding="utf-8-sig")

    has_markers = (
        bool(re.search(r"/\.\\(.+?)/\.\\", raw)) or           # 有标签标记
        bool(re.search(r"（【\{[^}]+\}】）", raw))               # 有占位符标记
    )

    if has_markers:
        # ── 标识符模板 ──────────────────────────────────────
        fields, labels = _extract_labels_and_fields(raw)
        print()
        print(c("── 标识符模板 ──────────────────────────────────────", "MAGENTA"))
        print(c(f"  字段: {', '.join(fields)}", "CYAN"))
        if any(labels.get(f, f) != f for f in fields):
            print(c("  标签→字段映射：", "CYAN"))
            for f in fields:
                lbl = labels.get(f, f)
                if lbl != f:
                    print(c(f"    /.\{lbl}/.\ → （【{{{f}}}】）", "YELLOW"))
        print()
        print(c("  请确认模板无误后继续。", "YELLOW"))
        ans = input(c("  按回车确认 / 输入 n 退出: ", "WHITE")).strip().lower()
        if ans in ("n", "no"):
            print(c("  [X] 用户取消", "RED"))
            return [], {}, ""
        return fields, labels, raw

    # ── 自然语言模板 ──────────────────────────────────────
    print()
    print(c("── 自然语言模板 ──────────────────────────────────────", "MAGENTA"))
    print(c("  未检测到 /.\ 标记，将自动转换为标识符模板。", "YELLOW"))

    identifier_raw = natural_to_identifier(raw)
    # 保存到模板文件同目录，加 _identifier 后缀
    id_path = path.parent / f"{path.stem}_identifier.docx"
    # 保存为真正的 docx 文件（不是纯文本）
    from docx import Document
    doc = Document()
    for line in identifier_raw.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    doc.save(str(id_path))

    fields, labels = _extract_labels_and_fields(identifier_raw)

    print()
    print(c(f"  已生成标识符模板: {id_path}", "GREEN"))
    print(c(f"  检测到字段: {', '.join(fields)}", "CYAN"))
    print()
    print(c("  转换后的模板内容：", "CYAN"))
    print(c("  " + "-" * 50, "WHITE"))
    for line in identifier_raw.split("\n"):
        print(c(f"  {line}", "WHITE"))
    print(c("  " + "-" * 50, "WHITE"))
    print()
    print(c("  请在文本编辑器中打开生成的标识符模板，修改满意后：", "YELLOW"))
    print(c(f"    {id_path}", "CYAN"))
    print()
    ans = input(c("  修改完成，按回车继续 / 输入 n 退出: ", "WHITE")).strip().lower()
    if ans in ("n", "no"):
        print(c("  [X] 用户取消", "RED"))
        return [], {}, ""

    # 重新读取用户可能修改过的标识符模板（真正的 docx）
    from docx import Document
    _doc = Document(str(id_path))
    identifier_raw = "\n".join(p.text for p in _doc.paragraphs)
    fields, labels = _extract_labels_and_fields(identifier_raw)
    print(c(f"  [OK] 确认，字段: {', '.join(fields)}", "GREEN"))
    return fields, labels, identifier_raw


def extract_all_fields(fields, video, transcript, labels=None):
    """
    从视频元数据和转写文本中提取所有字段。
    元数据 > LLM 语义提取 > 正则 fallback。
    提取不到留空。
    labels: dict {field_name: label_text} — /.\标签/.\ 文本用于 LLM 搜索目标
    """
    results = {}
    meta = video.copy()
    if labels is None:
        labels = {}

    for idx, field_name in enumerate(fields, 1):
        # 特殊字段：transcript -> 转写全文
        if field_name.lower() == "transcript":
            results[field_name] = transcript
            continue

        # 优先从元数据取
        if field_name in meta and meta[field_name]:
            results[field_name] = meta[field_name]
            print(c(f"    [{idx}/{len(fields)}] [OK] {field_name}: {str(meta[field_name])[:30]}", "GREEN"))
            continue

        # LLM 语义提取
        label = labels.get(field_name, field_name)
        print(c(f"    [{idx}/{len(fields)}] LLM 提取「{label}」...", "CYAN"))
        val = extract_field_with_llm(field_name, transcript, label=label, video_meta=video)
        if val:
            results[field_name] = val
            print(c(f"    [OK] {field_name}: {val[:30]}", "GREEN"))
            continue

        # 正则 fallback
        print(c(f"    [!] LLM 未提取到，尝试正则匹配...", "YELLOW"))
        val = extract_field_from_text(field_name, transcript)
        results[field_name] = val if val else ""
        if val:
            print(c(f"    [OK] {field_name}（正则）: {val[:30]}", "GREEN"))
        else:
            print(c(f"    [!] {field_name}: 未提取到内容", "YELLOW"))

    return results


def extract_field_from_text(field_name, text):
    """
    正则 fallback：从转写文本中按字面匹配提取字段值。
    只在 LLM 提取不到时使用。
    """
    if not text:
        return ""

    # 策略1：精确前缀（"辩题：XXX" 或 "辩题是 XXX"）
    patterns = [
        rf"{re.escape(field_name)}[：:]\s*([^\n]{{1,200}})",
        rf"{re.escape(field_name)}[是为]\s*([^\n]{{1,200}})",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            val = m.group(1).strip()
            val = re.sub(r"[。！？；]$", "", val)
            if val:
                return val[:300]

    # 策略2：字段名出现在文本中，取附近句子
    idx = text.find(field_name)
    if idx >= 0:
        start = max(0, idx)
        end = min(len(text), idx + 300)
        snippet = text[start:end].split("\n")
        if len(snippet) >= 2:
            return snippet[1].strip()[:300]
        return snippet[0].strip()[:300]

    return ""


# ═══════════════════════════════════════════════════
#  Section 5：生成 Word 文档
# ═══════════════════════════════════════════════════

def generate_output(fields, template_raw, all_results, output_path):
    """根据模板和提取结果生成 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for i, results in enumerate(all_results, 1):
        # 替换模板中的 （【{字段名}】） 占位符
        content = template_raw
        for field_name in fields:
            placeholder = f"（【{{{field_name}}}】）"
            val = results.get(field_name, "")
            content = content.replace(placeholder, str(val) if val else "（未提取到）")

        # 按行写入 Word
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # 去除 /.\保留字/.\ 标记（保留中间文字）
            display_line = re.sub(r"/\.\\(.+?)/\.\\", lambda m: m.group(1), line)

            p = doc.add_paragraph(display_line)
            # 标题样式（包含 = 或 多个 - ）
            if display_line.startswith("=") or display_line.startswith("-"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(14)

        # 每个视频之间加分页符（最后一个不加）
        if i < len(all_results):
            doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(c(f"  [OK] 已保存: {output_path.name}", "GREEN"))


# ═══════════════════════════════════════════════════
#  Section 6：主流程
# ═══════════════════════════════════════════════════

def main():
    # ── 交互模式（双击运行，无命令行参数）─────────────────────
    if len(sys.argv) <= 1:
        print(c("⚠️ 仅下载您拥有版权的原创内容，或取得授权的内容", "YELLOW"))
        print()
        print(c("=" * 52, "CYAN"))
        print(c("       From online video → script txt & answer docx", "CYAN"))
        print(c("=" * 52, "CYAN"))
        print()
        print(c("提示：直接回车使用默认值，输入路径可拖拽文件到窗口", "YELLOW"))
        print()

        # ═══════════════════════════════════════════════════
        #  第1步：必须依赖
        # ═══════════════════════════════════════════════════
        if not ensure_base_deps():
            input(c("\n按回车退出...", "YELLOW"))
            return

        # ═══════════════════════════════════════════════════
        #  第2步：Whisper 模型 + CUDA
        # ═══════════════════════════════════════════════════
        has_gpu, gpu_name, torch_cuda_ok, ct2_cuda_ok = detect_gpu()

        whisper_cuda = False
        whisper_model = None
        while whisper_model is None:
            print()
            print(c("── Whisper 模型选择 ──────────────────────────────", "MAGENTA"))
            print()
            WHISPER_OPTIONS = [
                ("tiny",   "~75MB",  "速度最快，中文准确度一般"),
                ("base",   "~145MB", "速度较快，准确度略高"),
                ("small",  "~488MB", "平衡之选，中文准度明显提升"),
                ("medium", "~1.5GB", "接近上限，中文效果很好（需6GB+ 显存）"),
            ]
            for size, sz, desc in WHISPER_OPTIONS:
                cached = _check_whisper_cache(size)
                tag = c(" [已缓存]", "GREEN") if cached else ""
                print(c(f"  {size:<6} {sz:<7} {desc}{tag}", "WHITE"))
            print()
            model_choice = input(c("  选择模型（默认 tiny）: ", "CYAN")).strip().lower()
            if not model_choice:
                model_choice = "tiny"
            elif model_choice not in ("tiny", "base", "small", "medium"):
                print(c(f"  [!] '{model_choice}' 不是有效选项，请重新选择", "YELLOW"))
                continue

            # 问 CUDA
            if has_gpu:
                print()
                print(c(f"  GPU: {gpu_name}", "GREEN"))
                if torch_cuda_ok:
                    print(c("  [*] torch(CUDA) 已安装", "GREEN"))
                while True:
                    ans = input(c("  是否启用 CUDA 加速？(y/n，默认 y): ", "CYAN")).strip().lower()
                    if ans in ("y", "yes", ""):
                        whisper_cuda = True
                        break
                    elif ans in ("n", "no", "fou", "bu"):
                        whisper_cuda = False
                        break
                    print(c("  [!] 无效输入，请输入 y 或 n", "YELLOW"))
            else:
                whisper_cuda = False

            if whisper_cuda:
                print(c("  -> Whisper: GPU 加速", "GREEN"))
            else:
                print(c("  -> Whisper: CPU", "YELLOW"))

            if ensure_whisper_setup(model_choice, whisper_cuda, gpu_name):
                whisper_model = model_choice

        # ═══════════════════════════════════════════════════
        #  第3步：LLM 模型 + llama.cpp CUDA
        # ═══════════════════════════════════════════════════
        llm_model = None
        llm_cuda = False
        while True:
            print()
            print(c("── LLM 模型选择 ──────────────────────────────────", "MAGENTA"))
            print()
            all_models = _all_llm_models()
            for key, info in all_models.items():
                status, _ = check_llm_model(key)
                if status == "ok":
                    tag = c(" [已下载]", "GREEN")
                elif status == "corrupt":
                    tag = c(" [损坏]", "RED")
                else:
                    tag = ""
                print(c(f"  {key:<22} {info['desc']}{tag}", "WHITE"))
            print()
            print(c("  输入模型名选择，空回车或 n = 不启用 LLM", "YELLOW"))
            llm_choice = input(c(f"  选择模型（默认 {LLM_DEFAULT}）: ", "CYAN")).strip().lower()
            if not llm_choice or llm_choice in ("n", "no"):
                print(c("  -> 不启用 LLM", "YELLOW"))
                break

            llm_model = _resolve_llm_key(llm_choice)
            if llm_model is None:
                continue

            # 问 llama.cpp CUDA
            if has_gpu:
                print()
                while True:
                    ans = input(c("  llama.cpp 是否用 CUDA 版？(y/n，默认 y): ", "CYAN")).strip().lower()
                    if ans in ("y", "yes", ""):
                        llm_cuda = True
                        break
                    elif ans in ("n", "no", "fou", "bu"):
                        llm_cuda = False
                        break
                    print(c("  [!] 无效输入，请输入 y 或 n", "YELLOW"))
            else:
                llm_cuda = False

            if llm_cuda:
                print(c("  -> llama.cpp: CUDA 版", "GREEN"))
            else:
                print(c("  -> llama.cpp: CPU 版", "YELLOW"))

            if ensure_llm_setup(llm_model, llm_cuda):
                break
            # 用户拒绝安装，回到选模型

        # ═══════════════════════════════════════════════════
        #  第4步：网址 + 模板
        # ═══════════════════════════════════════════════════
        print(c("\n── 参数输入 ──────────────────────────────────────", "MAGENTA"))
        print()

        url = None
        while not url:
            raw = input(c("  视频合集网址: ", "CYAN")).strip()
            url = extract_url(raw)
            if not url:
                print(c("  [!] 网址不能为空或格式无效，请粘贴包含 https:// 的链接", "YELLOW"))

        print()
        print(c("  模板文件（支持 .txt / .docx，直接写 标签名： 或 标签名 即可）：", "WHITE"))
        print(c("    标识符格式：/.\保留字/.\：（【{字段名}】）", "YELLOW"))
        print()
        template = None
        while not template:
            tmpl = input(c("  模板路径（必须提供，可拖拽文件）: ", "CYAN")).strip().strip('"')
            if not tmpl:
                print(c("  [!] 模板不能为空", "YELLOW"))
                continue
            if not Path(tmpl).exists():
                print(c(f"  [X] 文件不存在: {tmpl}", "RED"))
                continue
            template = tmpl

        print()
        seg = input(c("  时间段（回车=全部，支持 首3分 / 尾2分 / 1:30-5:45 / 0-3，逗号分隔）: ", "CYAN")).strip()
        segments = seg if seg else None

        lang = input(c("  语言（默认 zh，英文填 en）: ", "CYAN")).strip()
        language = lang if lang else "zh"

        # ═══════════════════════════════════════════════════
        #  第5步：模板确认 + 获取视频 + 跳过转写
        # ═══════════════════════════════════════════════════
        fields, labels, template_raw = confirm_template(template)
        if not fields or not template_raw:
            input(c("\n按回车退出...", "YELLOW"))
            return

        fix_ssl()
        videos = fetch_video_list(url, force_refresh=False)
        if not videos:
            input(c("\n按回车退出...", "YELLOW"))
            return

        skip_transcript = False

        # 检测已有转录稿
        existing = [v["bvid"] for v in videos
                    if (TRANSCRIPT_DIR / f"{v['bvid']}.txt").exists()]
        if existing:
            print()
            print(c(f"  检测到 {len(existing)}/{len(videos)} 个视频已有转录稿", "YELLOW"))
            print(c("    跳过转写 = 不跑 Whisper，但仍会跑繁转简 + LLM 处理", "YELLOW"))
            while True:
                ans = input(c("  是否跳过转写？(y/n，默认 n): ", "CYAN")).strip().lower()
                if ans in ("y", "yes"):
                    skip_transcript = True
                    break
                elif ans in ("n", "no", ""):
                    skip_transcript = False
                    break
                print(c("  [!] 无效输入，请输入 y 或 n", "YELLOW"))
            if skip_transcript:
                print(c("  -> 已有转录稿的视频将跳过转写", "GREEN"))
            else:
                print(c("  -> 将重新转写所有视频", "YELLOW"))

        # ═══════════════════════════════════════════════════
        #  加载 Whisper 模型
        # ═══════════════════════════════════════════════════
        device = "cuda" if whisper_cuda else "cpu"
        model = load_whisper_model(whisper_model, device)

        # ═══════════════════════════════════════════════════
        #  第6步：LLM 处理选项
        # ═══════════════════════════════════════════════════
        do_correct = False
        do_extract = False
        output = None
        if llm_model is not None:
            print()
            print(c("── LLM 处理选项 ─────────────────────────────────", "MAGENTA"))
            print()
            while True:
                ans = input(c("  是否进行同音字纠错？(y/n，默认 y): ", "CYAN")).strip().lower()
                if ans in ("y", "yes", ""):
                    do_correct = True
                    break
                elif ans in ("n", "no", "fou", "bu"):
                    do_correct = False
                    break
                print(c("  [!] 无效输入，请输入 y 或 n", "YELLOW"))
            if do_correct:
                print(c("  -> 将进行纠错，输出 {bvid}-纠错后.txt", "GREEN"))
            else:
                print(c("  -> 跳过纠错", "YELLOW"))

            print()
            while True:
                ans = input(c("  是否进行字段提取并生成 docx？(y/n，默认 y): ", "CYAN")).strip().lower()
                if ans in ("y", "yes", ""):
                    do_extract = True
                    break
                elif ans in ("n", "no", "fou", "bu"):
                    do_extract = False
                    break
                print(c("  [!] 无效输入，请输入 y 或 n", "YELLOW"))
            if do_extract:
                print(c("  -> 将提取字段并生成 docx", "GREEN"))
                print()
                default_out = f"output_{Path(template).stem}.docx"
                out = input(c(f"  输出文件名（默认 {default_out}）: ", "CYAN")).strip()
                output = out if out else default_out
            else:
                print(c("  -> 跳过提取", "YELLOW"))

        # ═══════════════════════════════════════════════════
        #  启动 LLM server（如需纠错或提取）
        # ═══════════════════════════════════════════════════
        server_proc = None
        llm_model_path = None
        if do_correct or do_extract:
            llm_status, llm_model_path = check_llm_model(llm_model)
            if llm_status != "ok":
                print(c("  [X] LLM 模型异常，无法进行纠错/提取", "RED"))
                do_correct = False
                do_extract = False
            else:
                server_proc = _start_llama_server(llm_model_path, llm_cuda)
                if not server_proc:
                    print(c("  [X] LLM server 启动失败", "RED"))
                    do_correct = False
                    do_extract = False

        # ═══════════════════════════════════════════════════
        #  处理每个视频
        # ═══════════════════════════════════════════════════
        AUDIO_DIR.mkdir(exist_ok=True)
        TRANSCRIPT_DIR.mkdir(exist_ok=True)
        OUTPUT_DIR.mkdir(exist_ok=True)

        all_results = []

        for video in videos:
            idx = video["index"]
            title = video["title"]
            bvid = video["bvid"]
            duration = video["duration"]

            print()
            print(c(f"  [{idx}/{len(videos)}] {title}", "CYAN"))
            print(c(f"  URL: https://www.bilibili.com/video/{bvid}", "YELLOW"))

            tr_path = TRANSCRIPT_DIR / f"{bvid}.txt"
            if skip_transcript and tr_path.exists():
                transcript = tr_path.read_text(encoding="utf-8")
                print(c(f"    [OK] 转录稿已存在，跳过转写 ({len(transcript)} 字)", "GREEN"))
            else:
                audio_path = download_audio(video, AUDIO_DIR)
                if not audio_path:
                    print(c(f"  [!] 跳过此视频（音频下载失败）", "YELLOW"))
                    continue
                seg_ranges = parse_segments(segments, duration)
                init_prompt = build_initial_prompt(labels, video_meta=video)
                transcript = transcribe_audio(audio_path, model, seg_ranges, language,
                                               initial_prompt=init_prompt)

            # 繁体→简体 + 清理
            transcript = fix_transcript_chinese(transcript)

            # 保存原始转写
            tr_path.write_text(transcript, encoding="utf-8")
            print(c(f"  [OK] 转写已保存: {tr_path.name}", "GREEN"))

            # LLM 纠错
            if do_correct and server_proc:
                transcript = correct_homophones_with_llm(
                    transcript, video, labels=labels, bvid=bvid
                )

            # LLM 提取
            if do_extract and server_proc:
                results = extract_all_fields(fields, video, transcript, labels=labels)
                all_results.append(results)

        # ═══════════════════════════════════════════════════
        #  生成输出
        # ═══════════════════════════════════════════════════
        if do_extract and all_results:
            if not output.lower().endswith(".docx"):
                output += ".docx"
            output_path = OUTPUT_DIR / output
            generate_output(fields, template_raw, all_results, output_path)
            print()
            print(c("=" * 52, "GREEN"))
            print(c(f"  [完成] 输出文件：{output_path}", "GREEN"))
            print(c("=" * 52, "GREEN"))
        elif do_extract and not all_results:
            print(c("[X] 没有成功处理任何视频。", "RED"))
        else:
            print()
            print(c("=" * 52, "GREEN"))
            n = len(videos)
            print(c(f"  [完成] 共处理 {n} 个视频", "GREEN"))
            print(c(f"  转写文本已保存至: {TRANSCRIPT_DIR}", "GREEN"))
            print(c("=" * 52, "GREEN"))

        # 关闭 server
        if server_proc:
            _stop_llama_server(server_proc)

        input(c("\n按回车退出...", "YELLOW"))
        return

    # ═══════════════════════════════════════════════════════
    #  命令行模式
    # ═══════════════════════════════════════════════════════
    parser = argparse.ArgumentParser(
        description="From online video → audio → script txt & answer docx ",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=TEMPLATE_HELP
    )
    parser.add_argument("--url",       help="合集/视频网址")
    parser.add_argument("--template",  default=None, help="模板文件路径（.txt / .docx）")
    parser.add_argument("--segments", help="时间段（如 首3分/尾2分,1:30-5:45,0-3）")
    parser.add_argument("--language",  default="zh",        help="语言（zh/en，默认 zh）")
    parser.add_argument("--model",     default="tiny",       help="Whisper 模型大小（tiny/base/small/medium）")
    parser.add_argument("--whisper-cuda", action="store_true", help="Whisper 启用 CUDA 加速")
    parser.add_argument("--llm-model", default=None,       help="LLM 模型（空 = 不启用）")
    parser.add_argument("--llm-cuda", action="store_true", help="llama.cpp 用 CUDA 版")
    parser.add_argument("--no-correct", action="store_true", help="跳过同音字纠错")
    parser.add_argument("--no-extract", action="store_true", help="跳过字段提取（不生成 docx）")
    parser.add_argument("--output",    help="输出 Word 文件名")
    parser.add_argument("--setup",    action="store_true", help="仅安装依赖，不处理视频")
    parser.add_argument("--refresh",  action="store_true", help="强制刷新视频列表缓存（重新获取）")
    parser.add_argument("--skip-transcript", action="store_true",
                        help="已有转录稿时跳过转写")
    parser.add_argument("--help-template", action="store_true", help="显示模板格式说明")
    args = parser.parse_args()

    if args.help_template:
        print(TEMPLATE_HELP)
        return

    if args.setup:
        ensure_base_deps()
        return

    if not args.url:
        print(c("[X] 请提供 --url 参数", "RED"))
        return

    if not args.template:
        print(c("[X] 请提供 --template 参数", "RED"))
        return

    args.language = args.language or "zh"
    args.model = args.model or "tiny"

    # ── GPU 检测 ──────────────────────────────────────
    has_gpu, gpu_name, torch_cuda_ok, ct2_cuda_ok = detect_gpu()

    whisper_cuda = args.whisper_cuda and has_gpu
    llm_cuda = args.llm_cuda and has_gpu

    # ── 依赖安装 ──────────────────────────────────────
    if not ensure_base_deps():
        return
    if not ensure_whisper_setup(args.model, whisper_cuda, gpu_name):
        print(c("[X] Whisper 依赖安装失败", "RED"))
        return

    llm_model = None
    if args.llm_model:
        llm_model = _resolve_llm_key(args.llm_model)
        if llm_model is None:
            print(c(f"[X] 无效的 --llm-model: '{args.llm_model}'", "RED"))
            return
        if not ensure_llm_setup(llm_model, llm_cuda):
            print(c("[X] LLM 依赖安装失败", "RED"))
            return

    # ── 模板确认 + 视频列表 ────────────────────────────
    fields, labels, template_raw = confirm_template(args.template)
    if not fields or not template_raw:
        return

    fix_ssl()
    videos = fetch_video_list(args.url, force_refresh=args.refresh)
    if not videos:
        return

    # ── 加载 Whisper ────────────────────────────────────
    device = "cuda" if whisper_cuda else "cpu"
    model = load_whisper_model(args.model, device)

    # ── LLM ────────────────────────────────────────────
    do_correct = llm_model is not None and not args.no_correct
    do_extract = llm_model is not None and not args.no_extract
    server_proc = None
    llm_model_path = None
    if do_correct or do_extract:
        llm_status, llm_model_path = check_llm_model(llm_model)
        if llm_status != "ok":
            print(c("  [X] LLM 模型异常", "RED"))
            do_correct = do_extract = False
        else:
            server_proc = _start_llama_server(llm_model_path, llm_cuda)
            if not server_proc:
                print(c("  [X] LLM server 启动失败", "RED"))
                do_correct = do_extract = False

    # ── 处理视频 ────────────────────────────────────────
    AUDIO_DIR.mkdir(exist_ok=True)
    TRANSCRIPT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)

    all_results = []

    for video in videos:
        idx = video["index"]
        title = video["title"]
        bvid = video["bvid"]
        duration = video["duration"]

        print()
        print(c(f"  [{idx}/{len(videos)}] {title}", "CYAN"))
        print(c(f"  URL: https://www.bilibili.com/video/{bvid}", "YELLOW"))

        tr_path = TRANSCRIPT_DIR / f"{bvid}.txt"
        if args.skip_transcript and tr_path.exists():
            transcript = tr_path.read_text(encoding="utf-8")
            print(c(f"    [OK] 转录稿已存在，跳过转写 ({len(transcript)} 字)", "GREEN"))
        else:
            audio_path = download_audio(video, AUDIO_DIR)
            if not audio_path:
                print(c(f"  [!] 跳过此视频（音频下载失败）", "YELLOW"))
                continue
            seg_ranges = parse_segments(args.segments, duration)
            init_prompt = build_initial_prompt(labels, video_meta=video)
            transcript = transcribe_audio(audio_path, model, seg_ranges, args.language,
                                           initial_prompt=init_prompt)

        transcript = fix_transcript_chinese(transcript)
        tr_path.write_text(transcript, encoding="utf-8")
        print(c(f"  [OK] 转写已保存: {tr_path.name}", "GREEN"))

        if do_correct and server_proc:
            transcript = correct_homophones_with_llm(
                transcript, video, labels=labels, bvid=bvid
            )

        if do_extract and server_proc:
            results = extract_all_fields(fields, video, transcript, labels=labels)
            all_results.append(results)

    # ── 生成输出 ────────────────────────────────────────
    if do_extract and all_results:
        output_name = args.output or f"output_{Path(args.template).stem}.docx"
        if not output_name.lower().endswith(".docx"):
            output_name += ".docx"
        output_path = OUTPUT_DIR / output_name
        generate_output(fields, template_raw, all_results, output_path)
        print()
        print(c("=" * 52, "GREEN"))
        print(c(f"  [完成] 输出文件：{output_path}", "GREEN"))
        print(c("=" * 52, "GREEN"))
    elif do_extract and not all_results:
        print(c("[X] 没有成功处理任何视频。", "RED"))
    else:
        print()
        print(c("=" * 52, "GREEN"))
        n = len(videos)
        print(c(f"  [完成] 共处理 {n} 个视频", "GREEN"))
        print(c(f"  转写文本已保存至: {TRANSCRIPT_DIR}", "GREEN"))
        print(c("=" * 52, "GREEN"))

    if server_proc:
        _stop_llama_server(server_proc)


if __name__ == "__main__":
    main()
